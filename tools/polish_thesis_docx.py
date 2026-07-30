#!/usr/bin/env python3
"""Hoàn thiện trang đầu, mục lục và danh mục viết tắt của khóa luận."""

from pathlib import Path
import os
import tempfile
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from update_thesis_docx import add_equation


ROOT = Path(__file__).resolve().parents[1]
PATH = ROOT / "thesis" / "khoa-luan.docx"


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


def raw_paragraph(text="", style=None, align=None, page_break=False):
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    if style:
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), style)
        ppr.append(ps)
    if align:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), align)
        ppr.append(jc)
    if page_break:
        ppr.append(OxmlElement("w:pageBreakBefore"))
    p.append(ppr)
    if text:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)
    return p


def toc_field():
    p = raw_paragraph()
    for kind, value in (
        ("begin", None),
        ("instr", ' TOC \\o "1-3" \\h \\z \\u '),
        ("separate", None),
        ("text", "Mục lục sẽ được cập nhật khi mở tài liệu."),
        ("end", None),
    ):
        r = OxmlElement("w:r")
        if kind == "instr":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = value
        elif kind == "text":
            node = OxmlElement("w:t")
            node.text = value
        else:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
            if kind == "begin":
                node.set(qn("w:dirty"), "true")
        r.append(node)
        p.append(r)
    return p


def abbreviation_table(doc):
    rows = [
        ("BLE", "Bluetooth Low Energy", "Bluetooth năng lượng thấp"),
        ("RSSI", "Received Signal Strength Indicator", "Chỉ số cường độ tín hiệu nhận"),
        ("HMAC", "Hash-based Message Authentication Code", "Mã xác thực dựa trên hàm băm"),
        ("MQTT", "Message Queuing Telemetry Transport", "Giao thức truyền dữ liệu IoT"),
        ("TLS", "Transport Layer Security", "Bảo mật lớp truyền"),
        ("OTA", "Over-the-Air", "Cập nhật chương trình qua mạng"),
        ("NVS", "Non-Volatile Storage", "Bộ nhớ không mất dữ liệu"),
        ("OOB", "Out-of-Band", "Kênh xác thực ngoài"),
        ("UWB", "Ultra-Wideband", "Băng thông siêu rộng"),
        ("UUID", "Universally Unique Identifier", "Mã định danh duy nhất"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(("Từ viết tắt", "Tên đầy đủ", "Ý nghĩa")):
        table.cell(0, i).text = value
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(0)
                if row_index == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
    return table


def insert_equation_after(doc, anchor, kind):
    p = add_equation(doc, kind)
    element = p._element
    element.getparent().remove(element)
    anchor._element.addnext(element)
    return p


def restore_equations(doc):
    """Khôi phục các công thức từng bị bộ chuyển đổi Word bỏ trống."""
    if any(node.tag == qn("m:oMathPara") for node in doc.element.iter()):
        return
    mappings = [
        ("ước lượng khoảng cách bằng log-distance path loss:", ["distance_simple"]),
        ("mô hình dùng là log-distance path loss", ["path_loss"]),
        ("Đảo công thức", ["distance_hat"]),
        ("Bộ lọc Kalman 1D coi", [
            "kalman_predict", "kalman_gain", "kalman_update", "kalman_cov"
        ]),
        ("HMAC [rfc2104] tạo mã xác thực", ["hmac"]),
        ("epoch_key dẫn xuất:", ["epoch_counter", "epoch_key"]),
    ]
    for marker, kinds in mappings:
        anchor = next(p for p in doc.paragraphs if marker in p.text)
        for kind in kinds:
            anchor = insert_equation_after(doc, anchor, kind)


def deduplicate_package(path):
    """Loại các ZIP entry trùng tên do một số trình sửa DOCX tạo ra."""
    with zipfile.ZipFile(path, "r") as source:
        latest = {}
        for info in source.infolist():
            latest[info.filename] = (info, source.read(info.filename))
    fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=path.parent)
    os.close(fd)
    temp = Path(temp_name)
    try:
        with zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
            for name, (info, data) in latest.items():
                target.writestr(info, data)
        temp.replace(path)
    finally:
        if temp.exists():
            temp.unlink()


def main():
    doc = Document(PATH)

    # Hai trang bìa.
    for p in [p for p in doc.paragraphs[:75] if p.text.strip() == "CAO TRỌNG PHƯỚC"]:
        p.text = "CAO TRỌNG PHƯỚC – TĂNG SĨ THÔNG"
    english_title = {
        "ROOM-LEVEL INDOOR POSITIONING SYSTEM",
        "BASED ON BLE MESH AND ESP32",
    }
    for p in list(doc.paragraphs[:80]):
        if p.text.strip() in english_title:
            remove_paragraph(p)
    for p in doc.paragraphs[:80]:
        if "HUỲNH HỮU THUẬN" in p.text and "TS." in p.text:
            p.text = "TS. HUỲNH HỮU THUẬN"
            following = p._element.getnext()
            following_text = "".join(following.itertext()) if following is not None else ""
            if "Trưởng Bộ môn Máy tính" not in following_text:
                p._element.addnext(raw_paragraph(
                    "Trưởng Bộ môn Máy tính – Hệ thống nhúng", align="center"
                ))
            break

    # Lời cam đoan và chữ ký.
    changes = {
        "Tôi xin cam đoan": "Chúng tôi xin cam đoan",
        "của bản thân tôi": "của chúng tôi",
        "do chính tôi thực hiện": "do chúng tôi thực hiện",
        "Tôi hoàn toàn chịu trách nhiệm": "Chúng tôi hoàn toàn chịu trách nhiệm",
    }
    for p in doc.paragraphs:
        text = p.text
        for old, new in changes.items():
            text = text.replace(old, new)
        if p.text.strip() == "Cao Trọng Phước":
            text = "Cao Trọng Phước – Tăng Sĩ Thông"
        if text != p.text:
            p.text = text

    # Tiêu đề và từ ngữ.
    for p in doc.paragraphs:
        text = p.text
        if text.startswith("Test "):
            text = "Kiểm thử " + text[5:]
        text = text.replace("(self-healing)", "").replace("end-to-end", "xuyên suốt")
        if text != p.text:
            p.text = text.strip()

    # Mục lục tự động và danh mục từ viết tắt.
    intro = next(p for p in doc.paragraphs if p.text.strip() == "GIỚI THIỆU")
    if not any(p.text.strip() == "MỤC LỤC" for p in doc.paragraphs):
        intro._element.addprevious(raw_paragraph(
            "MỤC LỤC", style="Heading1", align="center", page_break=True
        ))
        intro._element.addprevious(toc_field())
        intro._element.addprevious(raw_paragraph(
            "DANH MỤC TỪ VIẾT TẮT", style="Heading1", align="center", page_break=True
        ))
        table = abbreviation_table(doc)
        table_el = table._element
        table_el.getparent().remove(table_el)
        intro._element.addprevious(table_el)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    restore_equations(doc)
    doc.core_properties.author = "Cao Trọng Phước; Tăng Sĩ Thông"
    doc.save(PATH)
    deduplicate_package(PATH)


if __name__ == "__main__":
    main()
