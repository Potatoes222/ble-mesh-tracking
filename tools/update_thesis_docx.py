#!/usr/bin/env python3
"""Hoàn thiện khoa-luan.docx từ nội dung LaTeX và chuẩn hóa trình bày."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "thesis" / "khoa-luan.docx"
CHAPTERS = ROOT / "thesis" / "Content"


def delete_paragraph(p):
    el = p._element
    el.getparent().remove(el)
    p._p = p._element = None


def delete_table(table):
    el = table._element
    el.getparent().remove(el)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def clean_tex(text: str) -> str:
    replacements = {
        r"\%": "%",
        r"\_": "_",
        r"\&": "&",
        r"\$": "$",
        r"\{": "{",
        r"\}": "}",
        r"\times": "×",
        r"\rightarrow": "→",
        r"\geq": "≥",
        r"\leq": "≤",
        r"\pm": "±",
        r"\approx": "≈",
        r"\cdot": "·",
        "~": " ",
        "--": "–",
    }
    text = re.sub(r"\\(textbf|textit|texttt|emph)\{([^{}]*)\}", r"\2", text)
    text = re.sub(r"\\ref\{Chapter(\d)\}", r"Chương \1", text)
    text = re.sub(r"\\ref\{[^{}]+\}", "phần tương ứng", text)
    text = re.sub(r"\\cite\{[^{}]+\}", "", text)
    text = re.sub(r"\\label\{[^{}]+\}", "", text)
    text = re.sub(r"\\(begin|end)\{[^{}]+\}", "", text)
    text = re.sub(r"\\[a-zA-Z]+\{([^{}]*)\}", r"\1", text)
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("$", "")
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def add_body_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.add_run(clean_tex(text))
    return p


M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def m_el(tag: str):
    return OxmlElement(f"m:{tag}")


def m_run(text: str):
    r = m_el("r")
    rpr = m_el("rPr")
    sty = m_el("sty")
    sty.set(qn("m:val"), "p")
    rpr.append(sty)
    r.append(rpr)
    t = m_el("t")
    t.text = text
    r.append(t)
    return r


def m_group(*nodes):
    box = m_el("box")
    e = m_el("e")
    for node in nodes:
        e.append(node)
    box.append(e)
    return box


def m_frac(num, den):
    f = m_el("f")
    num_el = m_el("num")
    num_el.append(num)
    den_el = m_el("den")
    den_el.append(den)
    f.extend([num_el, den_el])
    return f


def m_sub(base, sub):
    obj = m_el("sSub")
    e = m_el("e")
    e.append(base)
    s = m_el("sub")
    s.append(sub)
    obj.extend([e, s])
    return obj


def m_sup(base, sup):
    obj = m_el("sSup")
    e = m_el("e")
    e.append(base)
    s = m_el("sup")
    s.append(sup)
    obj.extend([e, s])
    return obj


def m_hat(base):
    acc = m_el("acc")
    prop = m_el("accPr")
    char = m_el("chr")
    char.set(qn("m:val"), "\u0302")
    prop.append(char)
    e = m_el("e")
    e.append(base)
    acc.extend([prop, e])
    return acc


def m_delim(left: str, content, right: str):
    d = m_el("d")
    dpr = m_el("dPr")
    beg = m_el("begChr")
    beg.set(qn("m:val"), left)
    end = m_el("endChr")
    end.set(qn("m:val"), right)
    dpr.extend([beg, end])
    e = m_el("e")
    e.append(content)
    d.extend([dpr, e])
    return d


def sub(name: str, index: str):
    return m_sub(m_run(name), m_run(index))


def equation_nodes(kind: str):
    """Tạo cấu trúc công thức OMML cho các công thức dùng trong khóa luận."""
    if kind == "distance_simple":
        return [
            m_run("d = "),
            m_sup(
                m_run("10"),
                m_frac(m_group(sub("P", "0"), m_run(" − "), sub("P", "r")), m_run("10n")),
            ),
        ]
    if kind == "path_loss":
        inner = m_frac(m_run("d"), sub("d", "0"))
        return [
            sub("P", "r"), m_run("(d) = "), sub("P", "0"), m_run(" − 10n · "),
            m_sub(m_run("log"), m_run("10")), m_delim("(", inner, ")"),
            m_run(" + "), sub("X", "σ"),
        ]
    if kind == "distance_hat":
        exponent = m_frac(m_group(sub("P", "0"), m_run(" − "), sub("P", "r")), m_run("10n"))
        return [m_hat(m_run("d")), m_run(" = "), sub("d", "0"), m_run(" · "), m_sup(m_run("10"), exponent)]
    if kind == "kalman_predict":
        return [sub("p", "k"), m_run(" = "), sub("p", "k−1"), m_run(" + q")]
    if kind == "kalman_gain":
        return [sub("K", "k"), m_run(" = "), m_frac(sub("p", "k"), m_group(sub("p", "k"), m_run(" + r")))]
    if kind == "kalman_update":
        delta = m_group(sub("z", "k"), m_run(" − "), sub("x", "k−1"))
        return [sub("x", "k"), m_run(" = "), sub("x", "k−1"), m_run(" + "), sub("K", "k"), m_run(" · "), m_delim("(", delta, ")")]
    if kind == "kalman_cov":
        return [sub("p", "k"), m_run(" = "), m_delim("(", m_group(m_run("1 − "), sub("K", "k")), ")"), m_run(" · "), sub("p", "k")]
    if kind == "hmac":
        left = m_group(m_delim("(", m_group(m_run("K′ ⊕ opad")), ")"), m_run(" ∥ H"),
                       m_delim("(", m_group(m_delim("(", m_group(m_run("K′ ⊕ ipad")), ")"), m_run(" ∥ m")), ")"))
        return [m_run("HMAC(K, m) = H"), m_delim("(", left, ")")]
    if kind == "epoch_counter":
        frac = m_frac(m_run("esp_timer_get_time()"), m_run("1 giờ"))
        return [m_run("epoch_counter = ⌊"), frac, m_run("⌋")]
    if kind == "epoch_key":
        return [m_run("epoch_key = HKDF(master_key, epoch_counter)")]
    raise ValueError(kind)


def add_equation(doc: Document, kind: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    para = m_el("oMathPara")
    omath = m_el("oMath")
    for node in equation_nodes(kind):
        omath.append(node)
    para.append(omath)
    p._p.append(para)
    return p


def add_tex_equation(doc: Document, block: str):
    compact = re.sub(r"\s+", " ", block)
    if "d = 10^" in compact and "hat" not in compact:
        add_equation(doc, "distance_simple")
    elif "P_r(d)" in compact:
        add_equation(doc, "path_loss")
    elif r"\hat{d}" in compact:
        add_equation(doc, "distance_hat")
    elif "K_k &=" in compact:
        for kind in ("kalman_predict", "kalman_gain", "kalman_update", "kalman_cov"):
            add_equation(doc, kind)
    elif r"\text{HMAC}" in compact:
        add_equation(doc, "hmac")
    elif "epoch_counter" in compact or r"epoch\_counter" in compact:
        add_equation(doc, "epoch_counter")
        add_equation(doc, "epoch_key")


def parse_table(lines: list[str], start: int):
    rows = []
    caption = ""
    i = start
    while i < len(lines):
        line = lines[i].strip()
        cap = re.search(r"\\caption\{(.+)\}", line)
        if cap:
            caption = clean_tex(cap.group(1))
        if "&" in line and r"\\" in line:
            row = [clean_tex(x) for x in line.split(r"\\", 1)[0].split("&")]
            if row and not all(not x for x in row):
                rows.append(row)
        if line == r"\end{table}":
            return i, caption, rows
        i += 1
    return i, caption, rows


def add_tex(doc: Document, path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    para = []

    def flush():
        nonlocal para
        if para:
            text = clean_tex(" ".join(x.strip() for x in para))
            if text:
                add_body_paragraph(doc, text)
            para = []

    i = 0
    list_style = None
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line or line.startswith("%") or line.startswith(r"\label"):
            flush()
            i += 1
            continue
        m = re.match(r"\\chapter\{(.+)\}", line)
        if m:
            flush()
            p = doc.add_paragraph(style="Heading 1")
            p.add_run(clean_tex(m.group(1)))
            i += 1
            continue
        m = re.match(r"\\section\{(.+)\}", line)
        if m:
            flush()
            p = doc.add_paragraph(style="Heading 2")
            p.add_run(clean_tex(m.group(1)))
            i += 1
            continue
        m = re.match(r"\\subsection\{(.+)\}", line)
        if m:
            flush()
            p = doc.add_paragraph(style="Heading 3")
            p.add_run(clean_tex(m.group(1)))
            i += 1
            continue
        if line == r"\begin{itemize}":
            flush()
            list_style = "List Bullet"
            i += 1
            continue
        if line == r"\begin{enumerate}":
            flush()
            list_style = "List Number"
            i += 1
            continue
        if line in (r"\end{itemize}", r"\end{enumerate}"):
            flush()
            list_style = None
            i += 1
            continue
        if line.startswith(r"\item"):
            flush()
            add_body_paragraph(doc, line[5:].strip(), list_style or "List Bullet")
            i += 1
            continue
        if line.startswith(r"\begin{equation") or line.startswith(r"\begin{align"):
            flush()
            env = "align" if "align" in line else "equation"
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(rf"\end{{{env}"):
                block.append(lines[i])
                i += 1
            add_tex_equation(doc, "\n".join(block))
            i += 1
            continue
        if line == r"\begin{table}[htp]" or line.startswith(r"\begin{table}"):
            flush()
            end, caption, rows = parse_table(lines, i)
            if caption:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("Bảng. " + caption)
                r.bold = True
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for rr, values in enumerate(rows):
                    for cc in range(cols):
                        cell = table.cell(rr, cc)
                        cell.text = values[cc] if cc < len(values) else ""
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        set_cell_margins(cell)
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if rr == 0:
                                for run in p.runs:
                                    run.bold = True
                        if rr == 0:
                            set_cell_shading(cell, "D9EAF7")
            i = end + 1
            continue
        if line.startswith("\\begin{") or line.startswith("\\end{") or line in (r"\centering", r"\hline"):
            flush()
            i += 1
            continue
        if line.startswith(r"\[") or line.startswith(r"\]"):
            flush()
            i += 1
            continue
        para.append(line)
        i += 1
    flush()


def ensure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    settings = {
        "Heading 1": (16, WD_ALIGN_PARAGRAPH.CENTER, Pt(12), Pt(12)),
        "Heading 2": (14, WD_ALIGN_PARAGRAPH.LEFT, Pt(10), Pt(6)),
        "Heading 3": (13, WD_ALIGN_PARAGRAPH.LEFT, Pt(8), Pt(4)),
    }
    for name, (size, align, before, after) in settings.items():
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.alignment = align
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.space_before = before
        st.paragraph_format.space_after = after
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(13)
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        st.paragraph_format.space_after = Pt(0)

    if "Code" not in [s.name for s in styles]:
        code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Courier New"
        code.font.size = Pt(10)


def set_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def normalize_language(text: str) -> str:
    changes = {
        "Kèm theo phần chính là Tóm tắt (Việt + Anh)": "Kèm theo phần chính là Tóm tắt",
        "full stack": "hệ thống hoàn chỉnh",
        "end-to-end": "xuyên suốt",
        "End-to-end": "Xuyên suốt",
        "use case": "trường hợp sử dụng",
        "test ": "kiểm thử ",
        "Test ": "Kiểm thử ",
        "self-healing": "tự phục hồi",
        "server-side": "phía máy chủ",
        "re-flash": "nạp lại chương trình",
        "flash lại": "nạp lại",
        "scale": "quy mô",
        "fallback": "quay về bản đang hoạt động",
        "hardcode": "đặt cố định trong chương trình",
        "tự phục hồi (tự phục hồi)": "tự phục hồi",
    }
    for old, new in changes.items():
        text = text.replace(old, new)
    if text.startswith("p_k &="):
        text = (
            "Bước dự đoán: p(k) = p(k−1) + q. Hệ số Kalman: K(k) = p(k) / (p(k) + r). "
            "Giá trị lọc: x(k) = x(k−1) + K(k) × [z(k) − x(k−1)]. "
            "Sai số mới: p(k) = [1 − K(k)] × p(k)."
        )
    elif text.startswith("epoch_counter &="):
        text = (
            "Bộ đếm thời gian được tính theo từng giờ từ thời gian hoạt động của thiết bị. "
            "Khóa của mỗi khoảng thời gian được tạo bằng HKDF từ khóa chính và bộ đếm này."
        )
    return text


def main():
    doc = Document(DOCX)

    # Xóa bản tóm tắt tiếng Anh.
    start = end = None
    for idx, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "ABSTRACT":
            start = idx
        if start is not None and t == "GIỚI THIỆU":
            end = idx
            break
    if start is not None and end is not None:
        for p in list(doc.paragraphs[start:end]):
            delete_paragraph(p)

    # Thay toàn bộ các chương bằng nội dung LaTeX để không giữ ký hiệu chuyển đổi lỗi.
    cut = None
    for p in doc.paragraphs:
        if p.text.strip() == "GIỚI THIỆU":
            cut = p
            break
    if cut is not None:
        body = doc._element.body
        deleting = False
        for child in list(body):
            if child is cut._element:
                deleting = True
            if deleting and child.tag != qn("w:sectPr"):
                body.remove(child)

    for n in (1, 2, 3, 4, 5):
        add_tex(doc, CHAPTERS / f"chapter{n}.tex")

    # Chương kết luận trong nguồn mới chỉ có dàn ý, bổ sung nội dung rõ ràng.
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("KẾT LUẬN")
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Tóm tắt kết quả")
    add_body_paragraph(
        doc,
        "Khóa luận đã xây dựng một hệ thống xác định vị trí theo phòng gồm thẻ phát tín hiệu, "
        "ba bộ quét, nút chuyển tiếp, cổng kết nối và máy chủ ThingsBoard. Dữ liệu RSSI được "
        "thu tại bộ quét, chuyển qua BLE Mesh và gửi lên máy chủ bằng MQTT có mã hóa. "
        "Máy chủ chọn phòng bằng ngưỡng chênh lệch RSSI và hai lần xác nhận liên tiếp."
    )
    add_body_paragraph(
        doc,
        "Hệ thống có cơ chế xác thực gói quảng bá bằng HMAC, chống phát lại bằng số thứ tự, "
        "lưu cấu hình trong NVS và cập nhật chương trình qua HTTPS. Mã nguồn được chia theo "
        "từng chức năng nên có thể kiểm tra, sửa đổi và triển khai lại từng phần."
    )
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Hạn chế")
    for text in (
        "Hệ thống mới xác định được phòng, chưa xác định tọa độ cụ thể trong phòng.",
        "Kết quả RSSI còn chịu ảnh hưởng của tường, cửa, người di chuyển và nhiễu ở dải tần 2,4 GHz.",
        "Phạm vi thử nghiệm còn nhỏ, chưa đánh giá khi có nhiều hơn 10 thẻ hoặc nhiều khu vực độc lập.",
        "Thời gian sử dụng pin của thẻ chưa được theo dõi trong một đợt vận hành dài ngày.",
        "Hệ thống hiện phù hợp với một mạng nội bộ; việc quản lý nhiều địa điểm chưa được triển khai.",
    ):
        add_body_paragraph(doc, text, "List Bullet")
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("Hướng phát triển")
    for title, text in (
        ("Xác định tọa độ trong phòng", "Dùng ít nhất ba giá trị RSSI để ước lượng tọa độ và kết hợp bộ lọc hai chiều nhằm giảm dao động."),
        ("Hiệu chỉnh theo từng môi trường", "Đo dữ liệu tại từng phòng để chọn hệ số suy hao và ngưỡng chuyển phòng phù hợp."),
        ("Mở rộng quản lý thiết bị", "Bổ sung chức năng quản lý phiên bản, cấu hình và trạng thái của nhiều cổng kết nối ở nhiều địa điểm."),
        ("Tăng độ ổn định của đường truyền", "Đánh giá cách bố trí kênh Wi-Fi, vị trí nút chuyển tiếp và chế độ Bluetooth tầm xa trong môi trường có nhiều nhiễu."),
    ):
        p = doc.add_paragraph(style="Heading 3")
        p.add_run(title)
        add_body_paragraph(doc, text)

    ensure_styles(doc)

    # Khổ A4 và lề phù hợp cho bản in đóng gáy.
    for sec in doc.sections:
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(3.5)
        sec.right_margin = Cm(2.0)
        sec.header_distance = Cm(1.0)
        sec.footer_distance = Cm(1.0)
        set_page_number(sec)

    for p in doc.paragraphs:
        if p.style.name == "Normal":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if p.style.name == "Heading 1":
            p.paragraph_format.page_break_before = True
        old = p.text
        new = normalize_language(old)
        if new != old:
            for run in p.runs:
                run.text = ""
            p.add_run(new)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
                if r_idx == 0:
                    set_cell_shading(cell, "D9EAF7")
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True

    doc.core_properties.title = "Khóa luận tốt nghiệp – Hệ thống định vị trong nhà dựa trên BLE Mesh"
    doc.core_properties.subject = "BLE Mesh, ESP32 và ThingsBoard"
    doc.core_properties.author = "Cao Trọng Phước"
    doc.core_properties.keywords = "BLE Mesh, RSSI, ESP32, ThingsBoard, định vị trong nhà"
    doc.save(DOCX)


if __name__ == "__main__":
    main()
